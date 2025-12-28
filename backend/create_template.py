from docx import Document
from docx.shared import Pt
import os

def create_resume_template():
    doc = Document()
    
    # Title (Name)
    p = doc.add_paragraph()
    runner = p.add_run("{{ name }}")
    runner.bold = True
    runner.font.size = Pt(24)
    p.alignment = 1  # Center

    # Contact Info
    p = doc.add_paragraph()
    p.alignment = 1
    p.add_run("{{ email }} | {{ phone }} | {{ location }}")
    
    # Summary
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph("{{ summary }}")

    # Skills
    doc.add_heading("Skills", level=1)
    # comma separated list
    doc.add_paragraph("{{ skills | join(', ') }}")

    # Experience
    doc.add_heading("Experience", level=1)
    
    # Loop for experience
    p = doc.add_paragraph("{% for job in experience_list %}")
    
    p = doc.add_paragraph()
    runner = p.add_run("{{ job.title }}")
    runner.bold = True
    p.add_run(" at ")
    p.add_run("{{ job.company }}")
    p.add_run(" ({{ job.period }})")
    
    # Bullets loop within experience
    p = doc.add_paragraph("{% for b in job.bullets %}")
    doc.add_paragraph("- {{ b }}", style='List Bullet')
    doc.add_paragraph("{% endfor %}")
    
    doc.add_paragraph("{% endfor %}")

    # Projects
    doc.add_heading("Projects", level=1)
    p = doc.add_paragraph("{% for p in projects %}")
    
    p = doc.add_paragraph()
    runner = p.add_run("{{ p.name }}")
    runner.bold = True
    p.add_run(" - {{ p.tech_stack }}")
    
    doc.add_paragraph("{{ p.description }}")
    doc.add_paragraph("{% endfor %}")


    # Education
    doc.add_heading("Education", level=1)
    doc.add_paragraph("{% for edu in education %}")
    doc.add_paragraph("{{ edu.degree }}, {{ edu.institute }} ({{ edu.year }})")
    doc.add_paragraph("{% endfor %}")

    # Achievements
    doc.add_heading("Achievements", level=1)
    doc.add_paragraph("{% for ach in achievements %}")
    doc.add_paragraph("- {{ ach }}", style='List Bullet')
    doc.add_paragraph("{% endfor %}")

    output_path = os.path.join("templates", "resume_template.docx")
    doc.save(output_path)
    print(f"Template saved to {output_path}")

if __name__ == "__main__":
    create_resume_template()
